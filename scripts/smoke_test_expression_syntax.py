#!/usr/bin/env python3
"""
SMOKE TEST — RESOLVED 2026-08-26. Kept as a historical record only.

Confirmed via this script's own real API run: Doctavian DOES substitute
plain-text "{!fieldname}" placeholders from the uploaded data file. This
is now the basis for template_builder.py's implementation (see that
module's docstring for the full, confirmed syntax reference, including
Kanwal's direct confirmation and the mdoc:paragraph mechanism for
conditional content). No need to run this again for new work — it's
preserved here purely as the historical record of how the hypothesis
was tested and confirmed.

Original timeline (kept for context):
- 2026-08-24: hypothesized Doctavian's template engine reads PLAIN TEXT
  placeholders in its own expression syntax (e.g. "{!resource}", seen in
  their own "fieldExpression"/"variables" examples like "{!$now()}"),
  not native Word MERGEFIELD/IF field codes at all.
- 2026-08-25: Kanwal (Doctavian) reproduced our TEMPLATE_READ_FAILED
  error using our unmodified native-Word-field template and fixed it by
  changing only the uploaded data file (needed a top-level "data"
  wrapper key). This was read at the time as fully confirming the
  native-Word-field approach — but that reading was premature: it only
  proved the generate/download PIPELINE works, not that field
  substitution or the IF condition actually evaluate correctly.
- 2026-08-26: a real run (scripts/verify_doctavian_template.py) showed
  BOTH high- and low-risk documents came back with every merge field
  blank and identical static fallback text, with the field codes
  themselves completely untouched in the raw XML. This is exactly the
  failure mode this smoke test was originally designed to catch. It
  should be run for real now.

Run scripts/verify_doctavian_template.py FIRST — it now tests a
narrower, simpler hypothesis (put real values in the data file, keep
the native Word template unchanged) and is cheaper to rule out. Only
run this script if that one still fails (fields still blank / IF still
not evaluated even with real data present).

Usage:
    export DOCTAVIAN_API_KEY=...
    export DOCTAVIAN_ACCESS_TOKEN=...
    python scripts/smoke_test_expression_syntax.py
"""
import json
import os
import sys
import tempfile
import uuid
from pathlib import Path

from docx import Document

_AGENT_SRC = Path(__file__).resolve().parent.parent / "apps" / "agent" / "src"
sys.path.insert(0, str(_AGENT_SRC))

from tegata_agent.doctavian_client import DoctavianClient, DoctavianConfig  # noqa: E402


def main():
    api_key = os.environ.get("DOCTAVIAN_API_KEY")
    access_token = os.environ.get("DOCTAVIAN_ACCESS_TOKEN")
    base_url = os.environ.get("DOCTAVIAN_API_BASE_URL", "https://demo.api.doctavian.com")
    if not api_key or not access_token:
        print("Error: set DOCTAVIAN_API_KEY and DOCTAVIAN_ACCESS_TOKEN first.")
        sys.exit(1)

    client = DoctavianClient(
        DoctavianConfig(api_key=api_key, base_url=base_url, access_token=access_token)
    )

    # --- Build the simplest possible template: plain text, no Word fields ---
    doc = Document()
    doc.add_paragraph("Resource: {!resource}")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_docx_path = Path(tmp.name)
    doc.save(tmp_docx_path)
    print(f"Built minimal plain-text template at {tmp_docx_path}")

    print("\nUploading template...")
    uploaded_template = client.upload_template(tmp_docx_path)
    print(f"  {uploaded_template}")

    print("\nUploading data blob with a REAL value (not empty this time)...")
    with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as tmp:
        json.dump({"data": {"resource": "db_payment_prod"}}, tmp)
        tmp_data_path = Path(tmp.name)
    uploaded_data = client.upload_data(tmp_data_path)
    print(f"  {uploaded_data}")

    print("\nGenerating (plain-text {!resource} placeholder, no Word fields)...")
    try:
        result = client.generate_document(
            template_name="smoke-test-plain-text",
            template_urn=uploaded_template["id"],
            document_name=f"smoke-test-{uuid.uuid4().hex[:8]}",
            variables=[],
            external_request_id=f"smoke-test-{uuid.uuid4().hex[:8]}",
            data_urn=uploaded_data["id"],
        )
        print(f"\nGenerated: {result}")

        print("Downloading to check whether {!resource} was actually substituted...")
        docx_bytes = client.download_document(result["urn"])
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(docx_bytes)
            downloaded_path = Path(tmp.name)
        text = "\n".join(p.text for p in Document(downloaded_path).paragraphs if p.text.strip())
        downloaded_path.unlink(missing_ok=True)

        print(f"\nDownloaded document text:\n{text}\n")

        if "db_payment_prod" in text:
            print("=> CONFIRMED: plain-text {!...} placeholders DO get substituted.")
            print("=> template_builder.py needs a rewrite using plain-text placeholders")
            print("   instead of native Word MERGEFIELD/IF fields — see PROJECT_STATUS.md.")
        elif "{!resource}" in text:
            print("=> NOT confirmed: the literal '{!resource}' text passed through")
            print("   unsubstituted. Plain-text expression syntax isn't the answer either.")
            print("   Time to ask Kanwal directly what syntax their engine evaluates.")
        else:
            print("=> UNEXPECTED: neither the literal placeholder nor the real value")
            print("   appears in the output. Inspect the downloaded file manually.")
    except Exception as e:
        print(f"\nFAILED (generate or download raised): {e}")
        print("\n=> Inconclusive — the API call itself failed, so this doesn't tell us")
        print("   anything about placeholder substitution one way or the other.")

    tmp_docx_path.unlink(missing_ok=True)
    tmp_data_path.unlink(missing_ok=True)


if __name__ == "__main__":
    main()
