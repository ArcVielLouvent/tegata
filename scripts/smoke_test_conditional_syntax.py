#!/usr/bin/env python3
"""
Exploratory script — tests several CANDIDATE syntaxes for conditional /
branching logic inside Doctavian's plain-text expression templating
(confirmed 2026-08-26 via smoke_test_expression_syntax.py: Doctavian
reads plain-text "{!fieldname}" placeholders from the uploaded data
file's flat keys, and does NOT evaluate native Word MERGEFIELD/IF field
codes at all, regardless of whether real or empty data is supplied).

What's still unknown: the syntax for a CONDITIONAL expression (Tegata's
actual differentiator — the approval clause must read differently for
high vs low risk, not just substitute a flat value). We have no
documentation for this (their "Elements Reference" page is JS-rendered
and inaccessible to Claude's sandbox).

ROUND 1 RESULTS (2026-08-26, real API run):
    A. {!IF(required_approver_count == "2", "TWO approvers", "ONE approver")}
       -> rendered as EMPTY STRING (not literal passthrough)
    B. {!IF(required_approver_count = "2", "TWO approvers", "ONE approver")}
       -> rendered as EMPTY STRING
    C. {{#if required_approver_count == "2"}}TWO approvers{{else}}ONE approver{{/if}}
       -> rendered COMPLETELY UNCHANGED (literal passthrough)
    D. {!required_approver_count == "2" ? "TWO approvers" : "ONE approver"}
       -> rendered as EMPTY STRING

This is informative, not a dead end: A/B/D all used the "{!...}" wrapper
and came back EMPTY rather than passed through literally, meaning
Doctavian's "{!...}" parser DOES attempt to evaluate whatever's inside
it as a real expression (unlike "{{...}}", which it doesn't recognize
as a syntax at all and leaves completely untouched). The specific
function names/operators we guessed (bare "IF", "?:") just aren't
right.

ROUND 2 (this version): doctavian_client.py's own module docstring
documents a real example from Doctavian's docs — "{!$now()}" — note the
"$" prefix before the function name. This suggests built-in functions in
their expression language are namespaced with "$". Testing "$IF"/"$IIF"
variants on that basis.

Usage:
    export DOCTAVIAN_API_KEY=...
    export DOCTAVIAN_ACCESS_TOKEN=...
    python scripts/smoke_test_conditional_syntax.py

If NONE of round 2's candidates evaluate correctly either, STOP
guessing — this has now been tried across two well-reasoned rounds (7
candidates total) with concrete, specific results to hand Kanwal. Reply
to her thread with exactly this: "we've confirmed {!fieldname} plain-text
substitution works from the data file, and confirmed {!...} does attempt
to evaluate function-like expressions (unrecognized ones silently render
empty rather than erroring or passing through) — what is the correct
function name/syntax for an IF/conditional expression?"
"""
import json
import os
import sys
import tempfile
import uuid
from pathlib import Path

from docx import Document

_AGENT_SRC = Path(__file__).resolve().parent.parent / "apps" / "agent" / "src"
sys.path.insert(0, str(_AGENT_SRC))

from tegata_agent.doctavian_client import DoctavianClient, DoctavianConfig  # noqa: E402

# Round 1 candidates (A/B/C/D) all failed -- see docstring above for exact
# results. Round 2 tests the "$"-prefixed built-in function hypothesis.
CANDIDATES = {
    "E": 'E: {!$IF(required_approver_count == "2", "TWO approvers", "ONE approver")}',
    "F": 'F: {!$IF(required_approver_count = "2", "TWO approvers", "ONE approver")}',
    "G": 'G: {!$IIF(required_approver_count == "2", "TWO approvers", "ONE approver")}',
}


def main():
    api_key = os.environ.get("DOCTAVIAN_API_KEY")
    access_token = os.environ.get("DOCTAVIAN_ACCESS_TOKEN")
    base_url = os.environ.get("DOCTAVIAN_API_BASE_URL", "https://demo.api.doctavian.com")
    if not api_key or not access_token:
        print("Error: set DOCTAVIAN_API_KEY and DOCTAVIAN_ACCESS_TOKEN first.")
        sys.exit(1)

    client = DoctavianClient(
        DoctavianConfig(api_key=api_key, base_url=base_url, access_token=access_token)
    )

    doc = Document()
    doc.add_paragraph("Conditional syntax candidates, round 2 (required_approver_count = 2):")
    for label, text in CANDIDATES.items():
        doc.add_paragraph(text)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_docx_path = Path(tmp.name)
    doc.save(tmp_docx_path)
    print(f"Built candidate-testing template at {tmp_docx_path}")

    print("\nUploading template...")
    uploaded_template = client.upload_template(tmp_docx_path)
    print(f"  {uploaded_template}")

    print("\nUploading data (required_approver_count = 2, so all four candidates")
    print("should render 'TWO approvers' if their syntax is the correct one)...")
    with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as tmp:
        json.dump({"data": {"required_approver_count": "2"}}, tmp)
        tmp_data_path = Path(tmp.name)
    uploaded_data = client.upload_data(tmp_data_path)
    print(f"  {uploaded_data}")

    print("\nGenerating...")
    try:
        result = client.generate_document(
            template_name="smoke-test-conditional-syntax",
            template_urn=uploaded_template["id"],
            document_name=f"smoke-conditional-{uuid.uuid4().hex[:8]}",
            variables=[],
            external_request_id=f"smoke-conditional-{uuid.uuid4().hex[:8]}",
            data_urn=uploaded_data["id"],
        )
        print(f"Generated: {result}")

        print("\nDownloading...")
        docx_bytes = client.download_document(result["urn"])
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(docx_bytes)
            downloaded_path = Path(tmp.name)
        lines = [p.text for p in Document(downloaded_path).paragraphs if p.text.strip()]
        downloaded_path.unlink(missing_ok=True)

        print("\n" + "=" * 70)
        print("Downloaded document text:")
        for line in lines:
            print(f"  {line}")
        print("=" * 70)

        print("\nDIAGNOSIS (a candidate PASSED only if its line evaluated to")
        print("literally 'TWO approvers' with no leftover syntax characters):")
        any_passed = False
        for line in lines:
            if line.startswith(("E:", "F:", "G:")):
                label, content = line.split(":", 1)
                content = content.strip()
                if content == "TWO approvers":
                    print(f"  Candidate {label}: PASS — this is the correct syntax")
                    any_passed = True
                else:
                    print(f"  Candidate {label}: FAIL — rendered as: {content!r}")
        if not any_passed:
            print(
                "\nNone of round 2's candidates evaluated correctly either. Stop "
                "guessing here (7 candidates tried across 2 rounds) — reply to "
                "Kanwal's thread now with the specific, confirmed findings: "
                "plain-text {!field} substitution works, {!...} does attempt to "
                "evaluate function-like expressions (unrecognized ones render "
                "empty rather than erroring or passing through literally), but "
                "no IF/ternary/Handlebars syntax tried so far is recognized. "
                "Ask directly for the correct conditional-expression syntax."
            )
    except Exception as e:
        print(f"\nFAILED (generate or download raised): {e}")

    tmp_docx_path.unlink(missing_ok=True)
    tmp_data_path.unlink(missing_ok=True)


if __name__ == "__main__":
    main()
