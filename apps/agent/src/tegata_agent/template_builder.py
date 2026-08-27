"""
Tegata document template builder.

CONFIRMED SYNTAX (Kanwal, Doctavian, 2026-08-26 — no longer an
assumption): Doctavian's Document Generate engine does NOT evaluate
native Word MERGEFIELD/IF field codes at all — that was tried and
conclusively disproven (see PROJECT_STATUS.md's Phase 2 section for the
full investigation). Instead, Doctavian reads its own plain-text
"Elements"/"Expressions" templating language directly out of ordinary
paragraph text:

- Merge fields: {!fieldname} — substituted from the uploaded data
  file's top-level "data" object.
- Expressions (anything beyond a bare field — comparisons, functions,
  ternaries): {!$expression}, built on Jexl. Values are strings by
  default; wrap in toDecimal()/date()/etc. before doing typed
  operations. No IF()/IIF() function exists — use Jexl's native
  ternary (a == b ? x : y) for an inline value, or...
- ...for swapping entire blocks of content (Tegata's actual use case:
  the whole approval-clause sentence differs by risk tier, not just a
  word), use the mdoc:paragraph element: a whole paragraph is
  shown/hidden based on an expression in its "hidden" attribute. Two
  mdoc:paragraph blocks, each hidden under the opposite condition,
  implement an if/else.

Real example from Doctavian's own Mission 1 template (confirmed exact
syntax, including the quirk that the CLOSING tag repeats the "name"
attribute — not standard XML, but Doctavian's own convention, matched
here exactly, wrapped for readability in this docstring):

    <mdoc:paragraph name="exc10000"
        hidden="{!$toDecimal(sum(Customer.LineItems, "LineAmount")) < 10000}">
    Where the Total Contract Value exceeds $10,000, ...
    </mdoc:paragraph name="exc10000">

Each of the opening tag, the body text, and the closing tag is its own
separate paragraph in the .docx — not all inline in one paragraph.
That structure is mirrored exactly below.

No OOXML field-code manipulation (w:fldChar / w:instrText) is needed
anywhere anymore — every placeholder here is literal, ordinary paragraph
text. This is a complete rewrite of the pre-2026-08-26 version of this
file, which used native Word IF/MERGEFIELD fields; that version is
preserved in git history (see phase/2-doctavian commit history) for
reference, but should not be reused.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document


def build_tegata_template(output_path: str | Path) -> Path:
    """Builds the Tegata warrant .docx template with:
    - Plain {!fieldname} merge-field placeholders for always-visible
      fields (resource, reason, requester, risk score, tier, duration).
    - Two mdoc:paragraph elements implementing the approval clause's
      conditional branching (the "TWO approvers" vs "ONE approver"
      sentence), each hidden under the opposite condition on
      required_approver_count.

    Returns the path written.
    """
    output_path = Path(output_path)
    doc = Document()

    doc.add_heading("Tegata — Access Authorization Warrant", level=1)

    doc.add_paragraph("Resource: {!resource}")
    doc.add_paragraph("Requested by: {!requested_by}")
    doc.add_paragraph("Reason: {!reason}")
    doc.add_paragraph("Requested duration (minutes): {!requested_duration_minutes}")
    doc.add_paragraph("Approved maximum duration (minutes): {!max_duration_minutes}")
    doc.add_paragraph("Risk score: {!risk_score} / 100 (tier: {!risk_tier})")

    doc.add_heading("Approval Requirement", level=2)

    # Two mutually-exclusive mdoc:paragraph blocks implement the if/else:
    # required_approver_count is sent as a string (see warrant_variables.py),
    # so this compares against the string literal '2', matching how every
    # field value is treated as a string by default (see Expressions
    # Reference: "Every field value inside an expression is treated as a
    # string by default, regardless of its source type").
    doc.add_paragraph(
        '<mdoc:paragraph name="twoApprovers" hidden="{!$required_approver_count != \'2\'}">'
    )
    doc.add_paragraph("This request requires signatures from TWO approvers before it is valid.")
    doc.add_paragraph('</mdoc:paragraph name="twoApprovers">')

    doc.add_paragraph(
        '<mdoc:paragraph name="oneApprover" hidden="{!$required_approver_count == \'2\'}">'
    )
    doc.add_paragraph("This request requires a signature from ONE approver before it is valid.")
    doc.add_paragraph('</mdoc:paragraph name="oneApprover">')

    doc.add_paragraph(
        "This document is void if not signed within the approval window, "
        "and access is automatically revoked when the approved duration expires."
    )

    doc.save(output_path)
    return output_path
